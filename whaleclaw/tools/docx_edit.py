"""DOCX edit tool — targeted text replacement in an existing .docx file."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from whaleclaw.tools.base import Tool, ToolDefinition, ToolParameter, ToolResult
from whaleclaw.tools.deps import ensure_tool_dep


class DocxEditTool(Tool):
    """Edit an existing Word document by replacing exact text."""

    @property
    def definition(self) -> ToolDefinition:
        return ToolDefinition(
            name="docx_edit",
            description="修改现有 Word（.docx）文档中的文本，避免整份重写。",
            parameters=[
                ToolParameter(name="path", type="string", description="DOCX 文件绝对路径"),
                ToolParameter(name="old_text", type="string", description="要替换的原文"),
                ToolParameter(name="new_text", type="string", description="替换后的新文案"),
            ],
        )

    async def execute(self, **kwargs: Any) -> ToolResult:
        if not ensure_tool_dep("docx"):
            return ToolResult(success=False, output="", error="缺少依赖 python-docx")

        from docx import Document

        raw_path = str(kwargs.get("path", "")).strip()
        old_text = str(kwargs.get("old_text", "")).strip()
        new_text = str(kwargs.get("new_text", ""))

        if not raw_path:
            return ToolResult(success=False, output="", error="path 不能为空")
        if not old_text:
            return ToolResult(success=False, output="", error="old_text 不能为空")

        path = Path(raw_path).expanduser().resolve()
        if not path.is_file():
            return ToolResult(success=False, output="", error=f"文件不存在: {path}")
        if path.suffix.lower() != ".docx":
            return ToolResult(success=False, output="", error="仅支持 .docx 文件")

        try:
            doc = Document(str(path))
        except Exception as exc:
            return ToolResult(success=False, output="", error=f"DOCX 打开失败: {exc}")

        replaced = 0

        for para in doc.paragraphs:
            text = para.text or ""
            count = text.count(old_text)
            if count <= 0:
                continue
            para.text = text.replace(old_text, new_text)
            replaced += count

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text or ""
                    count = text.count(old_text)
                    if count <= 0:
                        continue
                    cell.text = text.replace(old_text, new_text)
                    replaced += count

        if replaced == 0:
            return ToolResult(success=False, output="", error="文档中未找到匹配文本")

        try:
            doc.save(str(path))
        except Exception as exc:
            return ToolResult(success=False, output="", error=f"DOCX 保存失败: {exc}")

        return ToolResult(success=True, output=f"已修改 {path}，替换 {replaced} 处文本")
